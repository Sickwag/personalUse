import os
import pandas as pd
from docx import Document
from docx.shared import Pt

def convert_excel_to_word(excel_file, output_dir):
    try:
        # 读取Excel文件，检测第一个工作表
        excel = pd.ExcelFile(excel_file)
        sheet_name = excel.sheet_names[0]
        df = pd.read_excel(excel_file, sheet_name=sheet_name, engine='openpyxl')
        
        # 检查必要的列是否存在
        required_columns = ['序号', '题干', 'A选项', 'B选项', 'C选项', 'D选项', '答案']
        if not all(column in df.columns for column in required_columns):
            print(f"警告: {excel_file} 中的列不完整。跳过此文件。")
            return
    except ValueError as e:
        print(f"错误: {excel_file} 中没有工作表或无法读取。错误信息: {e}。跳过此文件。")
        return
    except ImportError as e:
        print(f"错误: 缺少必要的库 {e}. 请安装所需的库。")
        return
    except Exception as e:
        print(f"错误: 无法读取 {excel_file}。错误信息: {e}")
        return

    # 获取Excel文件名（不含扩展名）
    base_name = os.path.splitext(os.path.basename(excel_file))[0]
    output_path = os.path.join(output_dir, f"{base_name}.docx")

    # 创建Word文档
    document = Document()

    # 遍历每一行数据
    for index, row in df.iterrows():
        # 获取各列数据
        try:
            title = f"{int(row['序号'])}. {row['题干']}"
        except (ValueError, TypeError):
            title = f"{row['序号']}. {row['题干']}"
        
        options = f"A. {row['A选项']}\tB. {row['B选项']}\tC. {row['C选项']}\tD. {row['D选项']}"
        answer = f"答案：{row['答案']}"

        # 添加内容到Word文档
        document.add_paragraph(title)
        document.add_paragraph(options)
        document.add_paragraph(answer)
        document.add_paragraph('')  # 添加一个空行作为分隔

    # 保存Word文档
    try:
        document.save(output_path)
        print(f"已生成: {output_path}")
    except Exception as e:
        print(f"错误: 无法保存 {output_path}。错误信息: {e}")

def process_folder(folder_path, output_dir):
    # 如果输出目录不存在，则创建
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 遍历文件夹中的所有文件
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith('.xlsx') or file.endswith('.xls'):
                excel_file = os.path.join(root, file)
                convert_excel_to_word(excel_file, output_dir)

def main():
    # 设置包含Excel文件的文件夹路径
    input_folder = r'D:\\cxdownload'  # 修改为你的文件夹路径

    # 设置输出Word文档的文件夹路径
    output_folder = r'D:\\cxdownload\\WordOutputs'  # 修改为你希望保存Word文档的文件夹路径

    process_folder(input_folder, output_folder)

if __name__ == "__main__":
    main()